from docx import Document
from db.workers import get_worker_by_id
from pathlib import Path
import os

def create_worker_docx(worker_id: int):
    """
    Создаёт анкету сотрудника в формате Word (.docx)
    по его id из таблицы Workers.
    """
    worker = get_worker_by_id(worker_id)
    if not worker:
        print(f"Сотрудник с id={worker_id} не найден.")
        return

    # Папка для хранения анкеты
    dump_dir = Path("dumps") / str(worker_id)
    dump_dir.mkdir(parents=True, exist_ok=True)
    file_path = dump_dir / f"{worker_id}.docx"

    doc = Document()
    doc.add_heading(f"Анкета сотрудника: {worker['fio']}", level=0)

    # Создаём таблицу с полями и значениями
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Поле"
    hdr_cells[1].text = "Значение"

    for field in ["fio", "position_id", "department_id", "phone", "email",
                  "birthdate", "address", "passport", "login"]:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = str(worker[field])

    doc.save(file_path)
    print(f"Анкета сохранена: {file_path}")
